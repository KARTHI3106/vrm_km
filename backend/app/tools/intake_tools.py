"""
Document Intake Agent tools — 8 tools for parsing, classifying, and extracting data.
"""
import io
import os
import re
import json
import logging
import hashlib
from datetime import datetime
from typing import Optional

import pdfplumber
from docx import Document as DocxDocument
import openpyxl
import pandas as pd
from langchain_core.tools import tool

from app.core.db import (
    create_document,
    update_document,
    check_duplicate_document,
    upload_file,
)
from app.core.redis_state import cache_get, cache_set
from app.core.llm import get_llm

MAX_TEXT_LENGTH = 50000
LLM_PREVIEW_LENGTH = 4000
LLM_DATE_PREVIEW_LENGTH = 3000

logger = logging.getLogger(__name__)

def _truncate_text(text: str, max_length: int = MAX_TEXT_LENGTH) -> str:
    """Intelligently truncate text, preserving start and end sections where signatures/summaries usually live."""
    if len(text) <= max_length:
        return text
    half = max_length // 2
    return text[:half] + f"\n\n...[TRUNCATED {len(text) - max_length} CHARACTERS]...\n\n" + text[-half:]

def validate_file(file_path: str, expected_type: Optional[str] = None) -> Optional[dict]:
    """Validates file existence, size, and magic bytes."""
    if not os.path.exists(file_path):
        return {
            "status": "error",
            "error_code": "FILE_NOT_FOUND",
            "error": f"File not found: {file_path}",
            "file_path": file_path
        }
    
    try:
        from app.config import get_settings
        settings = get_settings()
        upload_dir = os.path.abspath(settings.upload_dir)
        abs_path = os.path.abspath(file_path)
        if not abs_path.startswith(upload_dir):
            return {
                "status": "error",
                "error_code": "PATH_TRAVERSAL_DETECTED",
                "error": "File path is outside the allowed directory",
                "file_path": file_path
            }
    except Exception as e:
        logger.warning(f"Failed to check path traversal for {file_path}: {e}")

    file_size = os.path.getsize(file_path)
    MAX_SIZE = 50 * 1024 * 1024  # 50MB
    if file_size > MAX_SIZE:
        return {
            "status": "error",
            "error_code": "FILE_TOO_LARGE",
            "error": f"File size {file_size} exceeds maximum {MAX_SIZE}",
            "file_size": file_size,
            "max_size": MAX_SIZE
        }
        
    try:
        if expected_type:
            with open(file_path, 'rb') as f:
                header = f.read(4)
                if expected_type == "pdf" and header != b'%PDF':
                    return {
                        "status": "error",
                        "error_code": "INVALID_FILE_TYPE",
                        "error": "File is not a valid PDF",
                        "file_path": file_path
                    }
                elif expected_type in ("docx", "xlsx") and not header.startswith(b'PK\x03\x04'):
                    return {
                        "status": "error",
                        "error_code": "INVALID_FILE_TYPE",
                        "error": f"File is not a valid {expected_type.upper()}",
                        "file_path": file_path
                    }
    except PermissionError as e:
        return {"status": "error", "error_code": "PERMISSION_DENIED", "error": str(e), "file_path": file_path}
    except Exception as e:
        return {"status": "error", "error_code": "VALIDATION_ERROR", "error": str(e), "file_path": file_path}
        
    return None

def get_file_cache_key(file_path: str, tool_name: str) -> Optional[str]:
    """Generate a cache key based on file content hash, streamed in chunks to avoid RAM exhaustion."""
    try:
        h = hashlib.md5()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b""):
                h.update(chunk)
        return f"{tool_name}_{h.hexdigest()}"
    except Exception:
        return None


# ═══════════════════════════════════════════════════════════════════
# Tool 1: parse_pdf
# ═══════════════════════════════════════════════════════════════════

@tool
def parse_pdf(file_path: str) -> str:
    """
    Extract text and tables from a PDF document using pdfplumber.
    Returns the extracted text content along with any tables found.

    Args:
        file_path: Path to the PDF file on disk.
    """
    validation_err = validate_file(file_path, "pdf")
    if validation_err:
        return json.dumps(validation_err)
        
    cache_key = get_file_cache_key(file_path, "parse_pdf")
    if cache_key:
        cached = cache_get(cache_key)
        if cached:
            return json.dumps(cached)

    try:
        all_text = []
        all_tables = []

        with pdfplumber.open(file_path) as pdf:
            metadata = {
                "num_pages": len(pdf.pages),
                "metadata": pdf.metadata or {},
            }

            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                all_text.append(f"--- Page {i + 1} ---\n{text}")

                tables = page.extract_tables()
                if tables:
                    for t_idx, table in enumerate(tables):
                        all_tables.append(
                            f"[Table on page {i + 1}, #{t_idx + 1}]: "
                            + json.dumps(table)
                        )

        combined = "\n".join(all_text)
        if all_tables:
            combined += "\n\n=== TABLES ===\n" + "\n".join(all_tables)

        result_dict = {
            "status": "success",
            "text": _truncate_text(combined, MAX_TEXT_LENGTH),
            "num_pages": metadata["num_pages"],
            "has_tables": len(all_tables) > 0,
            "table_count": len(all_tables),
            "file_size": os.path.getsize(file_path),
            "file_path": file_path
        }
        if cache_key:
            cache_set(cache_key, result_dict, ttl=86400)
            
        result = json.dumps(result_dict)
        logger.info(f"Parsed PDF: {file_path} ({metadata['num_pages']} pages)")
        return result

    except PermissionError as e:
        logger.error(f"Permission denied for PDF {file_path}: {e}")
        return json.dumps({
            "status": "error",
            "error_code": "PERMISSION_DENIED",
            "error": f"Permission denied: {str(e)}",
            "file_path": file_path
        })
    except Exception as e:
        logger.error(f"Failed to parse PDF {file_path}: {e}, initiating automatic OCR fallback.")
        try:
            # automatic OCR fallback for parser failures
            return ocr_scan.invoke({"file_path": file_path})
        except Exception as ocr_error:
            logger.error(f"OCR fallback also failed for {file_path}: {ocr_error}")
            return json.dumps({
                "status": "error",
                "error_code": "PARSE_ERROR",
                "error": str(e),
                "error_type": type(e).__name__,
                "fallback_error": str(ocr_error),
                "file_path": file_path
            })


# ═══════════════════════════════════════════════════════════════════
# Tool 2: parse_docx
# ═══════════════════════════════════════════════════════════════════

@tool
def parse_docx(file_path: str) -> str:
    """
    Extract text and tables from a DOCX (Word) document.

    Args:
        file_path: Path to the DOCX file on disk.
    """
    validation_err = validate_file(file_path, "docx")
    if validation_err:
        return json.dumps(validation_err)
        
    cache_key = get_file_cache_key(file_path, "parse_docx")
    if cache_key:
        cached = cache_get(cache_key)
        if cached:
            return json.dumps(cached)

    try:
        doc = DocxDocument(file_path)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)

        tables_data = []
        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables_data.append(f"[Table #{t_idx + 1}]: {json.dumps(rows)}")

        combined = text
        if tables_data:
            combined += "\n\n=== TABLES ===\n" + "\n".join(tables_data)

        result_dict = {
            "status": "success",
            "text": _truncate_text(combined, MAX_TEXT_LENGTH),
            "num_paragraphs": len(paragraphs),
            "table_count": len(tables_data),
            "file_size": os.path.getsize(file_path),
            "file_path": file_path
        }
        if cache_key:
            cache_set(cache_key, result_dict, ttl=86400)

        result = json.dumps(result_dict)
        logger.info(f"Parsed DOCX: {file_path}")
        return result

    except Exception as e:
        logger.error(f"Failed to parse DOCX {file_path}: {e}")
        return json.dumps({
            "status": "error",
            "error_code": "PARSE_ERROR",
            "error": str(e),
            "error_type": type(e).__name__,
            "file_path": file_path
        })


# ═══════════════════════════════════════════════════════════════════
# Tool 3: parse_excel
# ═══════════════════════════════════════════════════════════════════

@tool
def parse_excel(file_path: str) -> str:
    """
    Extract data from all sheets of an Excel file.

    Args:
        file_path: Path to the Excel file on disk.
    """
    validation_err = validate_file(file_path, "xlsx")
    if validation_err:
        return json.dumps(validation_err)
        
    cache_key = get_file_cache_key(file_path, "parse_excel")
    if cache_key:
        cached = cache_get(cache_key)
        if cached:
            return json.dumps(cached)

    try:
        excel_data = {}
        xls = pd.ExcelFile(file_path)

        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            excel_data[sheet_name] = {
                "columns": df.columns.tolist(),
                "row_count": len(df),
                "data_preview": df.head(20).to_dict(orient="records"),
                "summary": df.describe(include="all").to_dict()
                if not df.empty
                else {},
            }

        result_dict = {
            "status": "success",
            "sheet_count": len(xls.sheet_names),
            "sheets": excel_data,
            "file_size": os.path.getsize(file_path),
            "file_path": file_path
        }
        
        if cache_key:
            cache_set(cache_key, result_dict, ttl=86400)

        result = json.dumps(result_dict, default=str)
        logger.info(
            f"Parsed Excel: {file_path} ({len(xls.sheet_names)} sheets)"
        )
        return result

    except Exception as e:
        logger.error(f"Failed to parse Excel {file_path}: {e}")
        return json.dumps({
            "status": "error",
            "error_code": "PARSE_ERROR",
            "error": str(e),
            "error_type": type(e).__name__,
            "file_path": file_path
        })


# ═══════════════════════════════════════════════════════════════════
# Tool 4: classify_document
# ═══════════════════════════════════════════════════════════════════

@tool
def classify_document(text: str) -> str:
    """
    Classify a document into a category using LLM analysis.
    Categories: SOC2, ISO27001, Insurance, DPA, Financial_Statements,
    BCP, Pen_Test_Report, Security_Questionnaire, Privacy_Policy, Other.

    Args:
        text: Extracted text content of the document (first 4000 chars recommended).
    """
    try:
        # Check cache (hash text content for caching)
        cache_key = f"classify_{hashlib.md5(text.encode()).hexdigest()}"
        cached = cache_get(cache_key)
        if cached:
            return json.dumps(cached)

        llm = get_llm()
        prompt = f"""You are a document classification expert for vendor risk assessment.

Classify the following document text into exactly ONE of these categories:
- SOC2 (SOC 2 Type 1 or Type 2 audit report)
- ISO27001 (ISO 27001 certification or audit)
- Insurance (General liability, cyber, E&O insurance certificates)
- DPA (Data Processing Agreement / Data Protection Agreement)
- Financial_Statements (Balance sheet, P&L, income statements)
- BCP (Business Continuity Plan / Disaster Recovery)
- Pen_Test_Report (Penetration testing report)
- Security_Questionnaire (Security assessment questionnaire responses)
- Privacy_Policy (Privacy policy document)
- Other (None of the above)

Respond in JSON format ONLY. Do not use blockquotes or markdown in the JSON string itself.
{{"classification": "<CATEGORY>", "confidence": <0.0-1.0>, "reasoning": "<brief explanation>"}}

Document text (first {LLM_PREVIEW_LENGTH} characters):
{text[:LLM_PREVIEW_LENGTH]}
"""
        response = llm.invoke(prompt)
        content = response.content if hasattr(response, "content") else str(response)

        # Try to extract JSON from response
        try:
            # Better JSON extraction handling markdown blocks
            clean_content = content.strip()
            if clean_content.startswith("```json"):
                clean_content = clean_content[7:-3].strip()
            elif clean_content.startswith("```"):
                clean_content = clean_content[3:-3].strip()
                
            json_match = re.search(r'\{[^{}]*\}', clean_content, re.DOTALL)
            if json_match:
                parsed = json.loads(json_match.group())
                
                result_dict = {
                    "status": "success",
                    "classification": parsed.get("classification", "Other"),
                    "confidence": parsed.get("confidence", 0.5),
                    "reasoning": parsed.get("reasoning", ""),
                }
                cache_set(cache_key, result_dict, ttl=86400)
                return json.dumps(result_dict)
        except json.JSONDecodeError:
            pass

        return json.dumps({
            "status": "success",
            "classification": "Other",
            "confidence": 0.3,
            "reasoning": "Could not parse LLM response, defaulting to Other",
        })

    except Exception as e:
        logger.error(f"Document classification failed: {e}")
        return json.dumps({"status": "error", "error_code": "LLM_CLASSIFICATION_ERROR", "error": str(e)})


# ═══════════════════════════════════════════════════════════════════
# Tool 5: extract_vendor_metadata
# ═══════════════════════════════════════════════════════════════════

@tool
def extract_vendor_metadata(text: str) -> str:
    """
    Extract structured vendor metadata from document text using LLM.
    Extracts: company name, address, contact, industry, employee count.

    Args:
        text: Extracted text from a vendor document (first 4000 chars recommended).
    """
    try:
        cache_key = f"metadata_{hashlib.md5(text.encode()).hexdigest()}"
        cached = cache_get(cache_key)
        if cached:
            return json.dumps(cached)
            
        llm = get_llm()
        prompt = f"""Extract vendor/company metadata from this document text.

Return JSON ONLY with these fields (use null for missing):
{{
    "company_name": "string",
    "address": "string",
    "contact_name": "string",
    "contact_email": "string",
    "phone": "string",
    "industry": "string",
    "employee_count": number or null,
    "website": "string",
    "domain": "string"
}}

Document text (first {LLM_PREVIEW_LENGTH} characters):
{text[:LLM_PREVIEW_LENGTH]}
"""
        response = llm.invoke(prompt)
        content = response.content if hasattr(response, "content") else str(response)

        try:
            clean_content = content.strip()
            if clean_content.startswith("```json"):
                clean_content = clean_content[7:-3].strip()
            elif clean_content.startswith("```"):
                clean_content = clean_content[3:-3].strip()

            # Use greedy match to capture nested JSON objects (e.g. nested address fields)
            json_match = re.search(r'\{.*\}', clean_content, re.DOTALL)
            if json_match:
                parsed = json.loads(json_match.group())
                result_dict = {"status": "success", "metadata": parsed}
                cache_set(cache_key, result_dict, ttl=86400)
                return json.dumps(result_dict)
        except json.JSONDecodeError:
            pass

        return json.dumps({
            "status": "success",
            "metadata": {},
            "note": "Could not extract metadata from LLM response",
        })

    except Exception as e:
        logger.error(f"Vendor metadata extraction failed: {e}")
        return json.dumps({"status": "error", "error_code": "LLM_METADATA_ERROR", "error": str(e)})


# ═══════════════════════════════════════════════════════════════════
# Tool 6: extract_dates
# ═══════════════════════════════════════════════════════════════════

@tool
def extract_dates(text: str) -> str:
    """
    Extract important dates from document text using regex and LLM.
    Finds: expiration dates, effective dates, issue dates, audit period dates.

    Args:
        text: Extracted document text to search for dates.
    """
    try:
        cache_key = f"dates_{hashlib.md5(text.encode()).hexdigest()}"
        cached = cache_get(cache_key)
        if cached:
            return json.dumps(cached)
            
        # Phase 1: Regex extraction for common date patterns
        date_patterns = [
            r'(?:expir(?:ation|es?|y)\s*(?:date)?[\s:]*)'
            r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
            r'(?:effective\s*(?:date)?[\s:]*)'
            r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
            r'(?:issue(?:d)?\s*(?:date)?[\s:]*)'
            r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
            r'(?:valid\s*(?:until|through|to)\s*[\s:]*)'
            r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
            # ISO format
            r'(\d{4}-\d{2}-\d{2})',
            # Written format
            r'(\w+\s+\d{1,2},?\s+\d{4})',
        ]

        regex_dates = []
        for pattern in date_patterns:
            matches = re.findall(pattern, text[:15000], re.IGNORECASE)
            regex_dates.extend(matches)

        # Phase 2: LLM extraction for contextual dates
        llm = get_llm()
        prompt = f"""Extract all important dates from this document text.

Return JSON ONLY:
{{
    "expiration_dates": ["YYYY-MM-DD", ...],
    "effective_dates": ["YYYY-MM-DD", ...],
    "issue_dates": ["YYYY-MM-DD", ...],
    "audit_period_start": "YYYY-MM-DD" or null,
    "audit_period_end": "YYYY-MM-DD" or null,
    "other_dates": [{{"label": "description", "date": "YYYY-MM-DD"}}]
}}

Document text (excerpt):
{text[:LLM_DATE_PREVIEW_LENGTH]}
"""
        response = llm.invoke(prompt)
        content = response.content if hasattr(response, "content") else str(response)

        llm_dates = {}
        try:
            clean_content = content.strip()
            if clean_content.startswith("```json"):
                clean_content = clean_content[7:-3].strip()
            elif clean_content.startswith("```"):
                clean_content = clean_content[3:-3].strip()
            json_match = re.search(r'\{.*\}', clean_content, re.DOTALL)
            if json_match:
                llm_dates = json.loads(json_match.group())
        except (json.JSONDecodeError, AttributeError):
            pass

        result_dict = {
            "status": "success",
            "regex_dates": regex_dates[:20],
            "llm_dates": llm_dates,
        }
        cache_set(cache_key, result_dict, ttl=86400)
        return json.dumps(result_dict)

    except Exception as e:
        logger.error(f"Date extraction failed: {e}")
        return json.dumps({"status": "error", "error_code": "LLM_DATE_ERROR", "error": str(e)})


# ═══════════════════════════════════════════════════════════════════
# Tool 7: store_document_metadata
# ═══════════════════════════════════════════════════════════════════

@tool
def store_document_metadata(
    vendor_id: str,
    file_name: str,
    file_type: str,
    classification: str,
    classification_confidence: float,
    extracted_text_summary: str,
    extracted_metadata: str,
    extracted_dates: str,
) -> str:
    """
    Store processed document metadata in the database.

    Args:
        vendor_id: The vendor UUID to associate the document with.
        file_name: Original file name.
        file_type: File extension (pdf, docx, xlsx).
        classification: Document classification category.
        classification_confidence: Classification confidence score (0-1).
        extracted_text_summary: Summary or first portion of extracted text.
        extracted_metadata: JSON string of extracted metadata.
        extracted_dates: JSON string of extracted dates.
    """
    try:
        # Check for duplicates
        if check_duplicate_document(vendor_id, file_name):
            return json.dumps({
                "status": "duplicate",
                "message": f"Document '{file_name}' already exists for this vendor.",
            })

        # Parse JSON strings
        try:
            meta = json.loads(extracted_metadata) if isinstance(extracted_metadata, str) else extracted_metadata
        except json.JSONDecodeError:
            meta = {"raw": extracted_metadata}

        try:
            dates = json.loads(extracted_dates) if isinstance(extracted_dates, str) else extracted_dates
        except json.JSONDecodeError:
            dates = {"raw": extracted_dates}

        doc_data = {
            "vendor_id": vendor_id,
            "file_name": file_name,
            "file_type": file_type,
            "classification": classification,
            "classification_confidence": classification_confidence,
            "extracted_text": extracted_text_summary[:10000],
            "extracted_metadata": meta,
            "extracted_dates": dates,
            "processing_status": "completed",
        }

        result = create_document(doc_data)
        return json.dumps({
            "status": "success",
            "document_id": result.get("id", ""),
            "message": f"Document '{file_name}' stored successfully.",
        })

    except Exception as e:
        logger.error(f"Failed to store document metadata: {e}")
        return json.dumps({"status": "error", "error": str(e)})


# ═══════════════════════════════════════════════════════════════════
# Tool 8: ocr_scan
# ═══════════════════════════════════════════════════════════════════

@tool
def ocr_scan(file_path: str) -> str:
    """
    Perform OCR on an image or scanned document using EasyOCR.
    Supports multi-language text extraction.

    Args:
        file_path: Path to the image file (PNG, JPG, TIFF, etc.).
    """
    validation_err = validate_file(file_path, None)
    if validation_err:
        return json.dumps(validation_err)
        
    cache_key = get_file_cache_key(file_path, "ocr_scan")
    if cache_key:
        cached = cache_get(cache_key)
        if cached:
            return json.dumps(cached)

    try:
        import easyocr

        reader = easyocr.Reader(["en"], gpu=False)
        results = reader.readtext(file_path)

        extracted_lines = []
        total_confidence = 0.0

        for bbox, text, confidence in results:
            extracted_lines.append(text)
            total_confidence += confidence

        avg_confidence = (
            total_confidence / len(results) if results else 0.0
        )
        full_text = "\n".join(extracted_lines)

        result_dict = {
            "status": "success",
            "text": _truncate_text(full_text, MAX_TEXT_LENGTH),
            "line_count": len(extracted_lines),
            "average_confidence": round(avg_confidence, 4),
            "quality": (
                "high" if avg_confidence > 0.8
                else "medium" if avg_confidence > 0.5
                else "low"
            ),
            "file_size": os.path.getsize(file_path),
            "file_path": file_path
        }
        
        if cache_key:
            cache_set(cache_key, result_dict, ttl=86400)
            
        return json.dumps(result_dict)

    except Exception as e:
        logger.error(f"OCR scan failed for {file_path}: {e}")
        return json.dumps({
            "status": "error", 
            "error_code": "OCR_PARSE_ERROR", 
            "error": str(e),
            "error_type": type(e).__name__,
            "file_path": file_path
        })


# ═══════════════════════════════════════════════════════════════════
# Collect all intake tools
# ═══════════════════════════════════════════════════════════════════

INTAKE_TOOLS = [
    parse_pdf,
    parse_docx,
    parse_excel,
    classify_document,
    extract_vendor_metadata,
    extract_dates,
    store_document_metadata,
    ocr_scan,
]

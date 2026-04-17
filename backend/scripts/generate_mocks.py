import os
from docx import Document
import pandas as pd

def create_mock_docs():
    output_dir = os.path.join(os.path.abspath(os.path.dirname(__file__)), '..', 'mock_documents')
    os.makedirs(output_dir, exist_ok=True)
    
    # 1. SOC 2 Type 2
    doc = Document()
    doc.add_heading('SOC 2 Type II Audit Report', 0)
    doc.add_paragraph('Company: Acme Corp Security Solutions\nAudit Period: Jan 1 2025 to Dec 31 2025')
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph('We have examined the controls of Acme Corp. The controls were suitably designed and operating effectively to provide reasonable assurance that the service organization\'s principal service commitments and system requirements were achieved based on the trust services criteria relevant to security, availability, and confidentiality.')
    doc.add_paragraph('There were no notable exceptions or deviations found during the testing period. Expiration Date: Dec 31 2026.')
    
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Control ID'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Status'
    row_cells = table.add_row().cells
    row_cells[0].text = 'CC1.1'
    row_cells[1].text = 'Access Control'
    row_cells[2].text = 'Passed'
    doc.save(os.path.join(output_dir, 'AcmeCorp_SOC2_Type2.docx'))
    
    # 2. DPA
    doc2 = Document()
    doc2.add_heading('Data Processing Agreement (DPA)', 0)
    doc2.add_paragraph('This Data Processing Agreement ("DPA") is effective as of January 15, 2026.')
    doc2.add_paragraph('Data Processor: Globex Corporation\nData Controller: The Customer.')
    doc2.add_paragraph('The Processor agrees to process personal data only on documented instructions from the Controller.')
    doc2.add_paragraph('Sub-processors: AWS (US-East), GCP (Europe).')
    doc2.save(os.path.join(output_dir, 'Globex_DPA.docx'))
    
    # 3. Pen Test
    doc3 = Document()
    doc3.add_heading('Penetration Test Executive Summary', 0)
    doc3.add_paragraph('Target: Initech Web Services\nDate of test: Feb 12, 2026')
    doc3.add_paragraph('Score: 8/10. Overall security posture is strong.')
    doc3.add_paragraph('Critical Findings: None.\nHigh Findings: 1 (Missing rate limiting on login endpoint).\nMedium Findings: 3.')
    doc3.add_paragraph('Recommendation: Immediate remediation of the High findings. The remaining medium findings are accepted risks.')
    doc3.save(os.path.join(output_dir, 'Initech_PenTest.docx'))

    # 4. Bad Pen Test (To test failure behavior)
    doc4 = Document()
    doc4.add_heading('Penetration Test Executive Summary', 0)
    doc4.add_paragraph('Target: Hooli Core Systems\nDate of test: March 1, 2026')
    doc4.add_paragraph('WARNING: Multiple critical vulnerabilities discovered.')
    doc4.add_paragraph('Critical Findings: 4 (SQL Injection, Remote Code Execution, Authentication Bypass).\nHigh Findings: 12.')
    doc4.add_paragraph('Recommendation: DO NOT DEPLOY. Immediate system shutdown recommended until critical vulnerabilities are patched.')
    doc4.save(os.path.join(output_dir, 'Hooli_Bad_PenTest.docx'))
    
    # 5. Financials (Excel)
    df = pd.DataFrame({
        'Metric': ['Total Revenue', 'Net Income', 'Operating Cash Flow', 'Total Assets', 'Total Liabilities'],
        'FY2025': [1200000, 150000, 300000, 5000000, 2000000],
        'FY2024': [1000000, 50000, 100000, 4500000, 2100000]
    })
    df.to_excel(os.path.join(output_dir, 'StarkIndustries_Financials.xlsx'), index=False)
    
    # 6. Bad Financials (Excel)
    df2 = pd.DataFrame({
        'Metric': ['Total Revenue', 'Net Income', 'Operating Cash Flow', 'Total Assets', 'Total Liabilities'],
        'FY2025': [50000, -1200000, -500000, 100000, 5000000],
        'FY2024': [100000, -800000, -200000, 500000, 3000000]
    })
    df2.to_excel(os.path.join(output_dir, 'MassiveDynamic_Bad_Financials.xlsx'), index=False)
    
    print(f"Mock documents successfully generated in: {os.path.abspath(output_dir)}")

if __name__ == '__main__':
    create_mock_docs()
